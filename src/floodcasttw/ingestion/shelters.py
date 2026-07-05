"""Shelter document parsing utilities."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document

from floodcasttw.io.csv_utils import write_csv

FIELDNAMES = ["鄉鎮市", "避難所名稱", "避難所地址", "避難所聯絡人", "收容人數", "來源檔案", "抽取時間"]
TITLE_MARKERS = ("清冊", "測試用", "避難收容處所")


def clean_contact(value: str) -> str:
    for label in ("主任", "總務", "校長", "里長", "村長"):
        value = value.replace(label, "")
    return value.strip()


def clean_capacity(value: str) -> str:
    digits = "".join(char for char in value if char.isdigit())
    return digits or "0"


def parse_shelter_line(line: str, source: str, extracted_at: str | None = None) -> dict[str, str] | None:
    text = line.strip()
    if not text or any(marker in text for marker in TITLE_MARKERS):
        return None

    tokens = text.split()
    if len(tokens) < 5:
        return None

    extracted_at = extracted_at or datetime.now().isoformat(timespec="seconds")
    return {
        "鄉鎮市": tokens[0],
        "避難所名稱": tokens[1],
        "避難所地址": " ".join(tokens[2:-2]),
        "避難所聯絡人": clean_contact(tokens[-2]),
        "收容人數": clean_capacity(tokens[-1]),
        "來源檔案": source,
        "抽取時間": extracted_at,
    }


def parse_table_row(cells: list[str], source: str, extracted_at: str) -> dict[str, str] | None:
    values = [cell.strip() for cell in cells if cell.strip()]
    if len(values) < 5 or any(marker in " ".join(values) for marker in TITLE_MARKERS):
        return None

    return {
        "鄉鎮市": values[0],
        "避難所名稱": values[1],
        "避難所地址": values[2],
        "避難所聯絡人": clean_contact(values[3]),
        "收容人數": clean_capacity(values[4]),
        "來源檔案": source,
        "抽取時間": extracted_at,
    }


def extract_from_docx(docx_path: Path) -> list[dict[str, str]]:
    if not docx_path.exists():
        raise FileNotFoundError(f"Input file not found: {docx_path}")

    document = Document(str(docx_path))
    extracted_at = datetime.now().isoformat(timespec="seconds")
    records: list[dict[str, str]] = []

    for table in document.tables:
        for row in table.rows:
            record = parse_table_row([cell.text for cell in row.cells], docx_path.name, extracted_at)
            if record:
                records.append(record)

    for paragraph in document.paragraphs:
        record = parse_shelter_line(paragraph.text, docx_path.name, extracted_at)
        if record:
            records.append(record)

    return records


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract shelter records from a DOCX file.")
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, default=Path("data/processed/shelters.csv"))
    args = parser.parse_args()

    records = extract_from_docx(args.input)
    count = write_csv(records, args.output, FIELDNAMES)
    print(f"[OK] Extracted {count} shelter records to {args.output}")


if __name__ == "__main__":
    main()
